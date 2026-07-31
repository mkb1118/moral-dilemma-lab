# -*- coding: utf-8 -*-
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Open the form
doc = Document(r'E:\我的桌面\班助选拔\软件学院2026级新生班助竞聘申请表.docx')
t = doc.tables[0]

# Helper: set cell text and apply font formatting
def set_cell_text(cell, text, font_name='宋体', font_size=Pt(11), bold=False, alignment=None):
    """Clear cell and set text with formatting."""
    # Remove existing paragraphs
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ''

    # Use first paragraph
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment

    # Clear existing runs
    for run in p.runs:
        run._element.getparent().remove(run._element)

    # Add new run
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = font_size
    run.font.bold = bold
    # Set East Asian font
    run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font_name)

def set_cell_text_multiline(cell, lines, font_name='宋体', font_size=Pt(11), bold=False):
    """Set multi-line text in a cell, one paragraph per line."""
    # Clear existing paragraphs - keep first, remove rest
    paras = cell.paragraphs
    # Remove all but first paragraph
    for p in list(paras)[1:]:
        p._element.getparent().remove(p._element)

    # Clear first paragraph
    first_p = paras[0]
    for run in list(first_p.runs):
        first_p._element.remove(run._element)

    for i, line in enumerate(lines):
        if i == 0:
            p = first_p
        else:
            p = cell.add_paragraph()

        # Set paragraph spacing
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.5

        run = p.add_run(line)
        run.font.name = font_name
        run.font.size = font_size
        run.font.bold = bold
        run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font_name)

# ==========================================
# Fill Basic Info (Rows 0-2)
# ==========================================

# R0C2: 姓名 → 马康博
set_cell_text(t.rows[0].cells[2], '马康博', font_size=Pt(12))
print('OK: 姓名 = 马康博')

# R0C4: 宿舍地址 → 2号楼216
set_cell_text(t.rows[0].cells[4], '2号楼216', font_size=Pt(12))
print('OK: 宿舍地址 = 2号楼216')

# R1C2: 性别 → 男
set_cell_text(t.rows[1].cells[2], '男', font_size=Pt(12))
print('OK: 性别 = 男')

# R1C4: 联系电话 → 19037360708
set_cell_text(t.rows[1].cells[4], '19037360708', font_size=Pt(12))
print('OK: 联系电话 = 19037360708')

# R2C2: 专业班级 → 软工融253
set_cell_text(t.rows[2].cells[2], '软工融253', font_size=Pt(12))
print('OK: 专业班级 = 软工融253')

# R2C4: QQ号码 → 1793461285
set_cell_text(t.rows[2].cells[4], '1793461285', font_size=Pt(12))
print('OK: QQ号码 = 1793461285')

# ==========================================
# Row 3: 现任职务 (C2 has gridSpan=3, merged C2-C4)
# ==========================================
set_cell_text(t.rows[3].cells[2], '学习委员', font_size=Pt(12))
print('OK: 现任职务 = 学习委员')

# ==========================================
# Row 4: 个人简历 (C1 has gridSpan=5, merged C1-C5)
# ==========================================
resume_lines = [
    '马康博，男，汉族，2007年11月出生，共青团员，河南商丘人。现就读于中原工学院软件学院软件工程（融合）专业，软工融253班，担任班级学习委员。',
    '',
    '主要经历：',
    '2025年9月至今，担任软工融253班学习委员，全面负责班级学习事务管理。每节课前检查多媒体设备运行情况，与教务处和媒体室保持密切沟通，及时处理设备故障，保障课堂教学顺利进行。负责各科作业的收集整理工作，逐一跟踪每位同学的作业提交情况，对未按时提交的同学耐心督促提醒，确保班级作业完成率保持较高水平。',
    '学习上勤奋刻苦，专业成绩良好，积极参加程序设计天梯赛等专业竞赛，不断提升编程能力和项目实践水平。主动组织班级互助学习小组，帮助学习上有困难的同学，分享学习笔记和期末复习方法，带动班级形成良好学习氛围。',
    '积极参与校园消防安全演练志愿服务活动，为校园安全贡献力量。承担音乐社活动摄影摄像工作，用镜头记录社团精彩瞬间。',
    '',
    '志愿服务经历：',
    '2026年1月至3月，在河南省商丘市民权县程庄镇人民政府参与社会实践和志愿服务，累计志愿时长100小时。服务期间积极参与社团日常事务协调、活动场地布置、现场秩序维护和服务信息统计等工作，得到了单位领导和同事的一致认可。',
]
set_cell_text_multiline(t.rows[4].cells[1], resume_lines, font_size=Pt(11))
print('OK: 个人简历 (共' + str(len(resume_lines)) + '段)')

# ==========================================
# Row 5: 奖惩情况 (C1 has gridSpan=5)
# ==========================================
awards_lines = [
    '获奖情况：',
    '1. 2025年天梯赛团体三等奖',
    '2. 2025年心理素质拓展运动会一等奖',
    '3. 2025年心理素质拓展运动会三等奖',
    '4. 2025年新生程序设计天梯赛三等奖',
    '5. 2026年校运动会三等奖',
    '6. 网络创业课程结业证书',
    '',
    '在校期间严格遵守校规校纪，无任何违纪处分记录。',
]
set_cell_text_multiline(t.rows[5].cells[1], awards_lines, font_size=Pt(11))
print('OK: 奖惩情况 (共' + str(len(awards_lines)) + '段)')

# ==========================================
# Row 6: 如何开展班助工作 (C1 has gridSpan=5)
# ==========================================
work_plan_lines = [
    '如果我有幸担任2026级新生班助，我将以饱满的热情和高度的责任感，从以下几个方面开展工作：',
    '',
    '一、协助学院做好新生迎新工作。提前熟悉新生名单和入学报到流程，在迎新现场热情接待新生及家长，协助办理各项入学手续，发放相关材料，耐心解答关于专业学习、校园生活等各类疑问，确保迎新工作高效有序进行。',
    '',
    '二、帮助新生尽快适应大学生活。结合自身从大一过来的亲身经历，向新生详细介绍校园环境、教学设施和生活服务资源，分享大学学习方法和时间管理技巧，帮助新生实现从高中到大学的平稳过渡。密切关注新生的心理状态和适应情况，及时发现并疏导存在适应困难的同学，引导他们积极融入大学生活。',
    '',
    '三、引导班级凝聚力和学风建设。组织新生破冰活动、主题班会和团建活动，增进同学之间的相互了解与友谊，营造团结互助、积极向上的班级氛围。发挥自身担任学习委员的经验优势，引导新生养成自主学习、独立思考的良好习惯，组织学习经验交流和期末复习帮扶活动，营造勤奋踏实的优良学风。',
    '',
    '四、协助处理班级日常事务。积极配合辅导员做好班级日常管理工作，及时准确传达学校和学院的各项通知要求，协助做好班委选举和团支部建设工作，指导班干部规范开展工作，逐步培养班级的自主管理能力。',
    '',
    '五、做好学院与班级之间的沟通桥梁。及时向新生传达学校和学院的各项政策规定和活动安排，同时主动收集同学们的合理诉求和意见建议，准确反馈给学院和辅导员，做到上传下达、信息畅通，促进师生之间的良性互动。',
    '',
    '六、协助新生教育、管理和服务工作。配合学院扎实开展新生入学教育、安全教育、心理健康教育和专业思想教育等各项活动，关注家庭困难、性格内向等特殊群体学生，主动关心帮助，协助解决实际困难，为新生健康成长保驾护航。',
    '',
    '我将以认真负责的态度和真诚服务的精神，全力以赴做好班助工作，努力成为新生的知心朋友和可靠引路人，不辜负学院和老师的信任与期望。',
]
set_cell_text_multiline(t.rows[6].cells[1], work_plan_lines, font_size=Pt(11))
print('OK: 如何开展班助工作 (共' + str(len(work_plan_lines)) + '段)')

# ==========================================
# Row 7: 学工办意见 - 留空，不填
# ==========================================
print('OK: 学工办意见 (留空)')

# ==========================================
# Save
# ==========================================
output_path = r'E:\我的桌面\班助选拔\软工融253-马康博-班助申请.docx'
doc.save(output_path)
print()
print('文件已保存: ' + output_path)
print('文件名格式: 班级+姓名+班助申请 → 软工融253-马康博-班助申请.docx')
