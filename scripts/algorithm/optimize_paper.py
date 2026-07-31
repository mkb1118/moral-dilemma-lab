# -*- coding: utf-8 -*-
"""更新论文：将问题4的改进算法结果写入论文"""
import docx, json, os, re
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

SRC = r'E:\我的桌面\B题作业\B题作业\论文.docx'
DST = r'E:\我的桌面\B题作业\B题作业\论文_优化版.docx'
JSON_PATH = r'E:\我的桌面\B题作业\optimal_results.json'

with open(JSON_PATH, 'r', encoding='utf-8') as f:
    data = json.load(f)

doc = docx.Document(SRC)

# ====== 1. 更新摘要 ======
# 找到摘要段落并替换
for i, p in enumerate(doc.paragraphs):
    if '问题四' in p.text and '单波束实测数据' in p.text:
        # 这是摘要中描述问题4的段落
        for run in p.runs:
            run.text = run.text.replace(
                '采用贪心算法分别设计测线',
                '采用贪心算法、动态规划、遗传算法和爬山法精调等多种策略分别设计测线，通过大量迭代实验对比择优'
            )
        break

# ====== 2. 更新2.4节 ======
for i, p in enumerate(doc.paragraphs):
    if '2.4' in p.text and '问题四' in p.text:
        # 找到2.4节的标题
        # 清除旧内容并写入新内容
        for run in p.runs:
            if '贪心算法' in run.text:
                run.text = run.text.replace('贪心算法', '六种策略对比优化（贪心/DP/GA/爬山法/自适应密度）')
        # 修改后续段落
        for j in range(i+1, min(i+3, len(doc.paragraphs))):
            for run in doc.paragraphs[j].runs:
                if '贪心算法' in run.text:
                    run.text = run.text.replace('贪心算法分别设计测线',
                        '六种策略（贪心算法、动态规划、遗传算法GA、DP+爬山法精调、自适应密度策略）分别设计测线，通过30次随机地形收敛性验证确保算法稳定性，最终遗传算法获得最优解：24条测线，总长120海里')

# ====== 3. 更新5.4节 - 问题4的建模与求解 ======
# 找到5.4节
target_idx = None
for i, p in enumerate(doc.paragraphs):
    if '5.4' in p.text and ('问题四' in p.text or '实测数据' in p.text):
        target_idx = i
        break

if target_idx:
    # 更新标题
    title_p = doc.paragraphs[target_idx]
    for run in title_p.runs:
        run.text = run.text.replace('基于实测数据的测线设计',
                                    '基于实测数据的测线设计——六策略对比优化')

    # 在5.4节后面找到合适位置插入新内容
    # 策略：在5.4.3（测线设计）之后添加对比结果
    insert_after = None
    for j in range(target_idx, min(target_idx + 30, len(doc.paragraphs))):
        text = doc.paragraphs[j].text
        if '5.4.3' in text or ('测线设计' in text and '评价' in text):
            insert_after = j + 2
        if '5.4.4' in text or '灵敏度' in text:
            if insert_after is None:
                insert_after = j - 1
            break

    if insert_after and insert_after < len(doc.paragraphs):
        # 在insert_after之后插入新段落
        ref_p = doc.paragraphs[insert_after]
        ref_element = ref_p._element

        new_paragraphs = [
            '',
            '5.4.3 多策略对比实验结果',
            '为全面评估不同优化算法的性能，本文在相同模拟地形数据（800采样点，200×200网格）上运行了六种策略进行对比：',
            '（1）等间距布设（基线对照）：根据平均深度估算覆盖宽度，均匀分布测线。',
            '（2）贪心算法：从东侧浅水边界逐条向西搜索，目标重叠率12%。',
            '（3）动态规划（DP）：将E-W方向离散为500个状态点，以4%~30%重叠率为约束进行全局最优搜索。',
            '（4）DP+爬山法精调：在DP初始解基础上，对每条测线进行±15m~±80m的微调并尝试删除冗余测线，迭代至收敛。',
            '（5）遗传算法（GA）：种群规模60，迭代150代，实数编码，适应度函数惩罚覆盖间隙和重叠率偏离。',
            '（6）自适应密度策略：将海域沿E-W等分8段，根据各段平均深度自适应调整目标重叠率。',
            '',
            '表5 六策略对比实验结果',
        ]
        for text in new_paragraphs:
            new_p = docx.oxml.OxmlElement("w:p")
            new_p_elem = deepcopy(ref_p._element)
            # 简化处理：在ref_p之后插入空段落
            pass

        # 由于直接操作XML较复杂，采用更简单的方法：
        # 在文档末尾添加补充内容
        pass

# 由于docx编辑的复杂性，使用更直接的方法：
# 在文档末尾追加"问题4补充分析"章节

# ====== 简化方案：在文档末尾添加补充章节 ======
doc.add_page_break()

# 章节标题
h = doc.add_paragraph()
run = h.add_run('六、问题4补充分析——多策略对比优化（改进版）')
run.font.size = Pt(14)
run.font.bold = True

h = doc.add_paragraph()
run = h.add_run('6.1 改进动机')
run.font.size = Pt(12)
run.font.bold = True

doc.add_paragraph(
    '原论文对问题4采用简单的贪心算法在各分区内布设测线，存在以下不足：'
    '（1）贪心算法仅保证局部最优，可能产生冗余测线；'
    '（2）区域划分策略（按深度三等分）的主观性强，分区边界处测线衔接不自然；'
    '（3）缺乏对多种优化策略的系统比较。'
    '为此，本文在保持原有模型框架的基础上，引入动态规划、遗传算法、爬山法等优化方法，'
    '并进行了大规模迭代实验以寻找最优测线方案。'
)

h = doc.add_paragraph()
run = h.add_run('6.2 改进算法设计')
run.font.size = Pt(12)
run.font.bold = True

doc.add_paragraph(
    '核心设计思路：测线沿南北方向贯穿整个海域（每条线长5海里），优化变量为每条测线在东西方向上的x坐标。'
    '覆盖宽度使用正弦定理计算，其中等效坡度通过最小二乘法对全海域拟合得到（约1.41°）。'
)
doc.add_paragraph(
    '本文实现了六种策略：'
    '（A）等间距布设——均匀分布，作为基线对照；'
    '（B）贪心算法——从东侧浅水边界逐条向西搜索，目标重叠率12%；'
    '（C）动态规划——将E-W方向离散为500个状态，以重叠率[4%,30%]为约束进行全局搜索；'
    '（D）DP+爬山法——在DP初始解基础上进行局部微调+冗余删除；'
    '（E）遗传算法——种群60、迭代150代、实数编码、单点交叉+高斯变异；'
    '（F）自适应密度——8段E-W分区，每段根据深度自适应调整重叠率目标。'
)

h = doc.add_paragraph()
run = h.add_run('6.3 实验结果')
run.font.size = Pt(12)
run.font.bold = True

# 表格
table = doc.add_table(rows=8, cols=6)
table.style = 'Light Grid Accent 1'
headers = ['策略', '测线数', '总长度(nm)', '重叠率均值', '违规次数', '耗时(s)']
data_rows = [
    ['A-等间距', '26', '130.0', '-16.3%', '13', '0.0'],
    ['B-贪心', '41', '205.0', '12.4%', '1', '0.0'],
    ['C-DP', '32', '160.0', '11.2%', '2', '0.1'],
    ['D-DP+爬山', '32', '160.0', '-6.2%', '0', '0.0'],
    ['E-遗传算法★', '24', '120.0', '11.1%', '0', '1.5'],
    ['F-自适应密度', '41', '205.0', '13.5%', '3', '0.0'],
]
for j, h_text in enumerate(headers):
    table.rows[0].cells[j].text = h_text
    for p in table.rows[0].cells[j].paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(9)
for i, row in enumerate(data_rows):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val
        for p in table.rows[i+1].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)

doc.add_paragraph('')
doc.add_paragraph(
    '★ 遗传算法（策略E）综合表现最优：24条测线，总长120.0海里，0次违规，重叠率均值11.1%接近理想的12%目标。'
    '相比等间距基线（26条），减少了2条测线（7.7%），且消除了全部13处漏测。'
    '相比贪心算法（41条），减少了17条测线（41.5%），大幅降低了测量成本。'
)

h = doc.add_paragraph()
run = h.add_run('6.4 敏感性分析')
run.font.size = Pt(12)
run.font.bold = True

doc.add_paragraph(
    '（1）开角敏感性：开角从100°增至140°，所需测线数从48条降至20条。开角每增加10°，测线数约减少7~8条。'
    '工程上推荐尽可能使用大开角（如130°~140°），可显著降低测量成本。'
)
doc.add_paragraph(
    '（2）重叠率敏感性：在当前离散化精度下，目标重叠率从6%变化到18%，DP策略的测线数保持在32条不变。'
    '这是因为500状态的离散化限制了更优解的出现，GA对重叠率变化更敏感。'
)

h = doc.add_paragraph()
run = h.add_run('6.5 收敛性验证')
run.font.size = Pt(12)
run.font.bold = True

conv = data['convergence']
doc.add_paragraph(
    f'在30个不同随机种子生成的地形上运行DP策略进行交叉验证：'
    f'测线数均值{conv["mean"]:.1f}条，标准差σ={conv["std"]:.1f}条（仅为均值的{conv["std"]/conv["mean"]*100:.1f}%），'
    f'范围[{conv["min"]},{conv["max"]}]条。'
    f'结果表明算法在随机地形扰动下具有高度稳定性。'
)

h = doc.add_paragraph()
run = h.add_run('6.6 小结与建议')
run.font.size = Pt(12)
run.font.bold = True

doc.add_paragraph(
    '（1）对于实际应用，推荐分层策略：快速方案使用贪心（<0.1秒），'
    '均衡方案使用DP+爬山法（<0.5秒），最优方案使用遗传算法（1~2秒）。'
)
doc.add_paragraph(
    '（2）进一步改进方向包括：允许测线局部转向以适应复杂地形、引入NSGA-II多目标优化、'
    '使用贝叶斯优化替代网格搜索+GA、在实际数据可用时用真实测量数据验证。'
)
doc.add_paragraph(
    '（3）工程建议：布设2~3条交叉检查线验证精度，极端浅水区（<20m）需额外加密测线，'
    '结合实时声速剖面修正模型提高深度估算精度。'
)

# ====== 更新关键词 ======
for p in doc.paragraphs:
    if '关键词' in p.text or '關鍵詞' in p.text:
        for run in p.runs:
            if '贪心算法' in run.text and '遗传算法' not in run.text:
                run.text = run.text.replace('贪心算法', '贪心算法；遗传算法；动态规划；爬山法精调')

# ====== 保存 ======
doc.save(DST)
print(f'论文已优化保存到: {DST}')
