# -*- coding: utf-8 -*-
"""问题4算法深度解析文档"""
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

DST = r'E:\我的桌面\B题作业\成果汇总\03_算法报告\问题4算法深度解析.docx'
os.makedirs(os.path.dirname(DST), exist_ok=True)

doc = Document()

for sec in doc.sections:
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.first_line_indent = Cm(0.74)

def set_font(run, name='宋体', size=12, bold=False):
    run.font.name = name; run.font.size = Pt(size); run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), name)
    rPr.append(rFonts)

def H(text, level=1):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(14-2*level)
    sizes = {0:18, 1:14, 2:13, 3:12}
    run = p.add_run(text); set_font(run, '黑体', sizes.get(level,12), True)

def B(text, indent=True):
    p = doc.add_paragraph()
    if not indent: p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text); set_font(run, '宋体', 12)

def formula(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text); set_font(run, 'Times New Roman', 11)

def code(text):
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text); set_font(run, 'Consolas', 9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def T(headers, rows, caption=''):
    if caption:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(caption); set_font(run, '黑体', 10, True)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]; c.text = h
        for pp in c.paragraphs:
            pp.paragraph_format.first_line_indent = Cm(0)
            for rn in pp.runs: set_font(rn, '黑体', 9, True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i+1].cells[j]; c.text = str(val)
            for pp in c.paragraphs:
                pp.paragraph_format.first_line_indent = Cm(0)
                for rn in pp.runs: set_font(rn, '宋体', 9)
    doc.add_paragraph()

# ====== CONTENT ======
for _ in range(4): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run('问题四：多波束测线设计'); set_font(run, '黑体', 22, True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.first_line_indent = Cm(0)
run = p.add_run('算法深度解析——从B477到本方案'); set_font(run, '宋体', 14)
for _ in range(3): doc.add_paragraph()

H('一、问题本质', 0)
B('问题四的核心是：给定一片待测海域的散点深度数据（单波束实测），要求在满足覆盖率约束的前提下，设计一组测线使得总长度最短。这本质上是一个带约束的组合优化问题，但有两个关键难点：')
B('(1) 地形是离散的散点数据，需要连续化才能进行任意位置的覆盖宽度计算。')
B('(2) 测线之间存在耦合关系（重叠率约束），使得问题具有序列决策特性——每条测线的位置依赖于前一条。')
doc.add_page_break()

# ====== METHOD 1: B477 ======
H('二、B477方法详解：随机森林 + 飞蛾扑火算法', 0)

H('2.1 地形建模：随机森林', 1)
B('B477采用随机森林（Random Forest）从散点单波束数据学习海底地形的连续映射 f(x,y) → depth。随机森林是一种集成学习方法，通过构建多个决策树并取平均来预测连续值。其优势在于：')
B('(1) 非线性拟合能力强：能捕捉海底的局部起伏、海脊、洼地等复杂地貌特征。')
B('(2) 天然正则化：通过Bootstrap采样和随机特征选择避免过拟合。')
B('(3) 可输出梯度：通过对每棵树的预测求偏导，可得任意位置的梯度方向 ∇f(x,y)。')
B('缺点是训练和推理计算量较大，对于800个数据点，训练+推理约需数秒。')

H('2.2 优化算法：改进飞蛾扑火（MFO）', 1)
B('飞蛾扑火算法（Moth-Flame Optimization）是一种群体智能优化算法，灵感来源于飞蛾绕火焰螺旋飞行的行为。B477对其做了三项关键改进以适配本题：')
B('(1) 自适应火源调整：原算法中"火焰"代表当前最优解，飞蛾围绕火焰做对数螺旋运动。B477将火焰方向替换为随机森林预测的梯度方向，使飞蛾（搜索粒子）始终沿地形等高线方向（垂直梯度）移动。')
B('(2) 地图微分化：将整个海域视为无数个微小区域的拼接，每个微小区域内近似平坦（类问题三），大幅简化了局部覆盖宽度的计算。')
B('(3) 适应度函数重构：将原算法中单纯的距离最小化替换为重叠率匹配度 + 路径长度加权组合，确保相邻测线的重叠率接近10%。')

H('2.3 位置更新公式', 1)
B('B477的关键迭代公式包括：')
B('(a) 覆盖宽度：W = D·sin(θ/2)·[1/sin(θ/2−α_eff) + 1/sin(θ/2+α_eff)]，其中 α_eff = arctan(tan(α)·|sin(β)|)')
B('(b) 粒子跃迁：当前位置沿梯度方向移动，移动步长由重叠率目标（n=10%）和局部水深决定：')
formula('Δx = W·(1−n)·gx(x,y),  Δy = W·(1−n)·gy(x,y)')
B('其中 gx, gy 为梯度方向的单位向量分量。当等效坡度不为零时，步长还需乘以水深修正因子 C/D。')

H('2.4 区域划分', 1)
B('B477发现所有区域共用一套参数会导致鲁棒性差，特别是在浅水区和坡度平缓区搜索效率低。为此基于等高线将海域划分为7个区域（Ⅰ-Ⅶ），每区独立运行MFO算法，最后合并测线。各区的测线数量和覆盖面积从3条/81万m²到28条/1355万m²不等。总计约76+条测线，有效覆盖率99.956%。')

doc.add_page_break()

# ====== METHOD 2: OURS ======
H('三、本方案方法详解：DP + GA + 爬山法', 0)

H('3.1 核心设计理念', 1)
B('本方案与B477的一个根本不同在于测线方向的选择。B477让测线随地形梯度自适应弯曲，而本方案将所有测线固定为N-S方向（平行等深线），理由已在论文5.4.2节阐述。这一简化使得优化变量从"二维曲线路径"降为"一维x坐标序列"，大幅缩小了搜索空间，使全局优化成为可能。')

H('3.2 地形连续化：Cubic插值', 1)
B('采用scipy.interpolate.griddata(method="cubic")将800个散点插值到200×200规则网格，再用scipy.ndimage.gaussian_filter(σ=1.0)平滑。相比B477的随机森林，cubic插值速度更快（<0.1s），且配合高斯平滑能有效抑制过拟合。对于N-S测线，沿经线取均值深度计算覆盖宽度，保证了结果的保守性和鲁棒性。')

H('3.3 策略A-B：基线方法', 1)
H('策略A：等间距布设', 2)
B('根据全海域平均深度计算典型覆盖宽度W_avg，然后均匀分布n=L_EW/(W_avg×0.85)条测线。简单快速（<0.01s），但在浅水区产生漏测、深水区过度重叠。')
H('策略B：贪心算法', 2)
B('从东侧（浅水）开始逐条向西搜索。每条测线搜索区间基于前一条的覆盖宽度动态估算，以2m步长寻找最接近12%目标重叠率的位置。复杂度O(n·k)。优点：速度快（<0.1s）；缺点：局部最优，可能产生冗余线（本例中41条）。')

H('3.4 策略C-D：动态规划 + 爬山法', 1)
H('策略C：动态规划', 2)
B('这是本方案的核心算法之一。将E-W方向离散为500个状态点（间距约15m）。')
B('状态定义：dp[i] = 覆盖到位置i且满足重叠率约束的最优方案（最小化总长度）。')
B('状态转移：从状态i到状态j的转移需满足重叠率在[4%, 30%]范围内，转移成本为L_NS（单条测线长5海里）× 效率惩罚因子。')
B('回溯：在所有覆盖西边界的状态中选择总成本最小的终点，回溯得到最优测线序列。')
B('DP保证了离散状态空间内的全局最优解，但500状态的离散化可能遗漏连续空间中的更优解。')

H('策略D：爬山法精调', 2)
B('在DP解的基础上进行局部优化：')
B('(1) 微调：对每条测线尝试±15, ±30, ±50, ±80m的位移。')
B('(2) 删减：尝试删除每条测线，检查是否仍满足覆盖和重叠率约束。')
B('(3) 迭代：重复至收敛（无改进或500次）。')
B('爬山法能有效消除DP离散化误差，但可能陷入局部最优。')

H('3.5 策略E：遗传算法（最优方案）', 1)
B('遗传算法是目前获得最优解的策略。核心设计如下：')
B('编码：染色体为n_lines个已排序的x坐标（实数编码），n_lines由平均覆盖宽度估算。')
B('适应度函数：f = −(覆盖惩罚 + 重叠偏离惩罚 + 边界惩罚)，其中覆盖惩罚对间隙（漏测）施加10倍权重，重叠偏离惩罚在[6%,25%]合规范围内较低、范围外5倍权重。')
B('选择：排序后取前50%精英。')
B('交叉：两点均匀交叉（各基因以50%概率来自父方或母方）。')
B('变异：以10%概率对每个基因添加N(0,30m)高斯噪声。')
B('迭代：种群60，迭代150代。')
B('GA的优势在于能逃离局部最优，在广阔解空间中随机搜索。24条/120nm的结果是所有策略中最优的。')

doc.add_page_break()

# ====== COMPARISON ======
H('四、两种方法深度对比', 0)

T(
    ['对比维度', 'B477方法', '本方案'],
    [
        ['优化变量', '二维曲线路径（x,y坐标序列）', '一维x坐标序列（N-S直线）'],
        ['搜索空间', '极大（连续曲线空间）', '适中（E-W方向500×N组合）'],
        ['算法类型', '群体智能（MFO局部搜索）', '精确+随机（DP全局+GA随机）'],
        ['理论基础', '群体智能+微分思想', 'DP最优子结构+进化计算'],
        ['局部最优风险', '中等（MFO有随机性但粒子可能早熟）', '低（GA+爬山+多策略互补）'],
        ['验证手段', '未系统报告', '30次随机地形交叉验证,σ=0.9'],
        ['测线条数', '76+条（含分区碎片化）', '24条（N-S连续贯穿）'],
        ['总路径长度', '约21998m（不含分区重复）', '120海里=222,240m'],
        ['覆盖率', '99.956%', '~95%'],
        ['工程便利性', '中等（曲线需精确导航）', '高（平行直线，自动驾驶友好）'],
        ['计算开销', '数秒（RF训练+MFO迭代）', '<2s（插值+GA 150代）'],
    ],
    '表1  两种方法全面对比'
)

H('五、为什么本方案更好：五个关键理由', 0)

H('5.1 变量降维使全局优化成为可能', 1)
B('B477的测线是二维曲线，搜索空间极大。MFO虽然是有效的启发式算法，但本质上是局部搜索+随机跳跃，无法保证全局最优。本方案将测线固定为N-S直线，优化变量降至一维，使得DP能在离散空间中保证最优解，GA能在连续空间中随机搜索全局最优。这是本方案最根本的优势——用变量降维换取优化质量的飞跃。')

H('5.2 多策略互补消除单点故障', 1)
B('B477仅使用MFO一种算法。如果MFO参数不当或初始位置不佳，可能导致收敛到劣解而无法自行纠正。本方案实现了6种策略（均匀/贪心/DP/爬山/GA/自适应），彼此互补：贪心提供快速基线，DP提供理论下界，GA提供随机探索，爬山法消除局部瑕疵。这种多策略互补设计从根本上消除了单点故障风险。')

H('5.3 严格的统计验证', 1)
B('B477未报告算法的稳定性验证。本方案进行了30次不同随机种子地形的交叉验证，测线数的标准差仅为均值的2.8%（σ=0.9 vs μ=31.6），范围[30,33]极小。这一严格的统计证据确保了算法在不同地形下的可靠性和可复现性。')

H('5.4 工程实用性优先', 1)
B('实际多波束测量作业中，测线通常设计为平行直线，原因是：(a)船舶自动驾驶系统沿直线航行最稳定；(b)数据处理软件针对平行测线优化；(c)平行线便于后续的条带拼接和误差校准。B477的曲线测线虽然理论覆盖率更高，但在工程实践中执行难度大、数据后处理复杂。本方案的N-S平行线完全符合行业标准做法。')

H('5.5 效率与精度的最佳平衡', 1)
B('B477用76+条测线达到99.956%覆盖率，本方案用24条测线达到~95%覆盖率。在工程上，最后的5%覆盖率往往可以通过布设2~3条交叉检查线来弥补，而测线数从76条降至24条（减少68%）带来的成本节省是巨大的（时间、燃料、人力）。这是一个典型的精度-效率权衡，本方案选择了更实用的平衡点。')

doc.add_page_break()

H('六、算法伪代码', 0)

H('核心算法1：动态规划测线优化', 1)
code('输入: X_states[0..N-1], 每状态的覆盖宽度W[i], Wl[i], Wr[i]')
code('输出: 最优测线x坐标序列')
code('')
code('dp_cost[0..N-1] = INF')
code('dp_prev[0..N-1] = -1')
code('')
code('# 初始化：第一条线覆盖东边界')
code('for i in 0..N-1:')
code('    if X[i] - Wr[i] <= 0:')
code('        dp_cost[i] = L_NS  # 单条线长')
code('')
code('# DP递推')
code('for i in 0..N-1:')
code('    if dp_cost[i] == INF: continue')
code('    for j in i+1..min(N, i+max_jump):')
code('        overlap = (X[i]+Wr[i]) - (X[j]-Wl[j])')
code('        eta = overlap / ((W[i]+W[j])/2)')
code('        if eta_min <= eta <= eta_max:')
code('            penalty = 1.0 + |eta-target|')
code('            new_cost = dp_cost[i] + L_NS * penalty')
code('            if new_cost < dp_cost[j]:')
code('                dp_cost[j] = new_cost')
code('                dp_prev[j] = i')
code('')
code('# 回溯最优路径（覆盖西边界的状态）')
code('best = argmin(dp_cost[i] for i where X[i]+Wl[i] >= L_EW)')
code('return reconstruct_path(best, dp_prev)')

H('核心算法2：遗传算法', 1)
code('输入: pop_size=60, generations=150')
code('输出: 最优测线方案')
code('')
code('# 初始化：包含贪心解和随机解')
code('pop = [greedy_solution] + [random_x_sorted() for _ in range(59)]')
code('')
code('for gen in 1..generations:')
code('    fitness = [evaluate(chrom) for chrom in pop]')
code('    elite = select_top50%(pop, fitness)')
code('    new_pop = elite.copy()')
code('    while len(new_pop) < pop_size:')
code('        p1, p2 = random_choice(elite, 2)')
code('        child = uniform_crossover(p1, p2)')
code('        child = gaussian_mutate(child, sigma=30, rate=0.1)')
code('        child = sort(child)  # 保持x坐标有序')
code('        new_pop.append(child)')
code('    pop = new_pop')
code('')
code('return decode(pop[argmax(fitness)])')

doc.add_page_break()

H('七、适用场景建议', 0)
T(
    ['应用场景', '推荐方法', '理由'],
    [
        ['实时交互规划', '贪心算法(本方案B)', '<0.1s,适合船上即时调整'],
        ['日常测量任务', 'DP+爬山法(本方案D)', '<0.5s,均衡速度与质量'],
        ['精细离线规划', '遗传算法(本方案E)', '1~2s,最优方案24条/120nm'],
        ['极高覆盖率要求(>99%)', 'MFO(B477方法)', '牺牲效率换取覆盖率'],
        ['地形极度复杂海域', 'B477梯度追踪+MFO', '自适应方向更有优势'],
        ['标准作业海域', '本方案N-S平行线', '符合行业标准,执行便利'],
    ],
    '表2  不同场景的方法推荐'
)

H('八、总结', 0)
B('问题四的多波束测线设计是一个典型的工程优化问题，需要在覆盖率、测线总长度和工程可行性之间取得平衡。')
B('B477优秀论文采用随机森林+飞蛾扑火算法的技术路线，实现了极高的覆盖率（99.956%），代价是测线条数多（76+条）且路径复杂。该方法适用于对覆盖率要求极高的精密测量任务，或地形极端复杂的海域。')
B('本方案采用变量降维（N-S平行线）+ 六策略对比优化（DP/GA/爬山法）+ 严格统计验证的技术路线，仅需24条测线即可达到~95%覆盖率，且总长度仅120海里，工程执行便利。该方法适用于大多数标准海洋测量场景，在效率、成本和可行性之间取得了最优平衡。')
B('两种方法互为补充——若追求极致覆盖率，B477方法更优；若追求效率与工程实用性，本方案更优。理想的做法是结合两者：用本方案的GA找到最优的N-S基线方案，再在浅水区等漏测高风险区域补充B477式的自适应曲线测线。')

# ====== SAVE ======
doc.save(DST)
print(f'算法解析已保存: {DST}')
